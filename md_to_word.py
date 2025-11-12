#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
将 Markdown 文件转换为 Word 文档
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import re

def convert_md_to_word(md_file, output_file):
    """将 Markdown 文件转换为 Word 文档"""

    # 读取 Markdown 文件
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()

    # 创建 Word 文档
    doc = Document()

    # 设置文档默认字体
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    doc.styles['Normal'].font.size = Pt(11)

    # 按行处理
    lines = md_content.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]

        # 处理标题
        if line.startswith('# '):
            # 一级标题
            heading = line.replace('# ', '')
            p = doc.add_heading(heading, level=1)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('## '):
            # 二级标题
            heading = line.replace('## ', '')
            p = doc.add_heading(heading, level=2)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('### '):
            # 三级标题
            heading = line.replace('### ', '')
            p = doc.add_heading(heading, level=3)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif line.startswith('#### '):
            # 四级标题
            heading = line.replace('#### ', '')
            p = doc.add_heading(heading, level=4)
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT

        # 处理分隔线
        elif line.strip() == '---':
            doc.add_paragraph('_' * 50)

        # 处理表格
        elif '|' in line and i + 1 < len(lines) and '|' in lines[i + 1]:
            # 收集表格行
            table_lines = []
            while i < len(lines) and '|' in lines[i]:
                if not lines[i].strip().startswith('|--'):  # 跳过表格分隔行
                    table_lines.append(lines[i])
                i += 1
            i -= 1

            # 解析表格
            if table_lines:
                # 解析表头
                headers = [cell.strip() for cell in table_lines[0].split('|')[1:-1]]

                # 创建表格
                table = doc.add_table(rows=1, cols=len(headers))
                table.style = 'Light Grid Accent 1'

                # 填充表头
                hdr_cells = table.rows[0].cells
                for j, header in enumerate(headers):
                    hdr_cells[j].text = header
                    # 表头加粗
                    for paragraph in hdr_cells[j].paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                # 填充数据行
                for row_line in table_lines[1:]:
                    cells_data = [cell.strip() for cell in row_line.split('|')[1:-1]]
                    row_cells = table.add_row().cells
                    for j, cell_data in enumerate(cells_data):
                        # 移除 Markdown 的 HTML 标签
                        cell_data = re.sub(r'<br>', '\n', cell_data)
                        row_cells[j].text = cell_data

        # 处理无序列表
        elif line.strip().startswith('- ') or line.strip().startswith('* '):
            text = line.strip()[2:]
            # 移除 Markdown 标记
            text = remove_md_formatting(text)
            p = doc.add_paragraph(text, style='List Bullet')

        # 处理有序列表
        elif re.match(r'^\d+\.\s', line.strip()):
            text = re.sub(r'^\d+\.\s', '', line.strip())
            text = remove_md_formatting(text)
            p = doc.add_paragraph(text, style='List Number')

        # 处理复选框列表
        elif line.strip().startswith('- [ ]') or line.strip().startswith('- [x]'):
            checked = '[x]' in line
            text = re.sub(r'- \[[ x]\]\s*', '', line.strip())
            text = remove_md_formatting(text)
            checkbox = '☑ ' if checked else '☐ '
            p = doc.add_paragraph(checkbox + text)

        # 处理普通段落
        elif line.strip():
            # 跳过 Markdown 链接定义
            if not re.match(r'^\[.+\]:\s*#', line):
                text = remove_md_formatting(line)
                p = doc.add_paragraph(text)

        # 处理空行
        else:
            # 在非表格、非列表情况下添加空行
            if i > 0 and lines[i-1].strip():
                doc.add_paragraph()

        i += 1

    # 保存文档
    doc.save(output_file)
    return output_file

def remove_md_formatting(text):
    """移除 Markdown 格式标记"""
    # 移除粗体
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    # 移除斜体
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    # 移除行内代码
    text = re.sub(r'`(.+?)`', r'\1', text)
    # 移除链接 [text](url)
    text = re.sub(r'\[(.+?)\]\(.+?\)', r'\1', text)
    # 移除 emoji 标记(如果需要)
    # text = re.sub(r':[a-z_]+:', '', text)
    return text

if __name__ == '__main__':
    import sys
    import os

    # 设置标准输出编码
    if sys.platform == 'win32':
        import io
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

    md_file = r'c:\Users\clark\OneDrive\桌面\vans\SERVICE_AGREEMENT_TEMPLATE.md'
    output_file = r'c:\Users\clark\OneDrive\桌面\vans\SERVICE_AGREEMENT_TEMPLATE.docx'

    try:
        result = convert_md_to_word(md_file, output_file)
        print("Conversion successful!")
        print(f"Word document saved to: SERVICE_AGREEMENT_TEMPLATE.docx")
    except Exception as e:
        print(f"Conversion failed: {e}")
        import traceback
        traceback.print_exc()
        print("\nPlease make sure python-docx is installed:")
        print("pip install python-docx")
